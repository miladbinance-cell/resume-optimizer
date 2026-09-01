import os

def extract_text_from_file(filepath: str, ext: str) -> str:
    """
    Extract text from uploaded resume file.
    Supports PDF, DOCX, and plain text.
    """
    ext = ext.lower()

    try:
        if ext == ".pdf":
            return _extract_from_pdf(filepath)
        elif ext == ".docx":
            return _extract_from_docx(filepath)
        elif ext in (".txt", ".md", ".rtf"):
            return _extract_from_text(filepath)
        else:
            # Try all methods as fallback
            for method in [_extract_from_pdf, _extract_from_docx, _extract_from_text]:
                try:
                    text = method(filepath)
                    if text.strip():
                        return text
                except Exception:
                    continue
            return ""
    except Exception as e:
        return f"Could not extract text: {str(e)}"


def _extract_from_pdf(filepath: str) -> str:
    """Extract text from PDF using PyPDF2"""
    from PyPDF2 import PdfReader
    reader = PdfReader(filepath)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()


def _extract_from_docx(filepath: str) -> str:
    """Extract text from DOCX using python-docx"""
    from docx import Document
    doc = Document(filepath)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + " | "
            text += "\n"
    return text.strip()


def _extract_from_text(filepath: str) -> str:
    """Extract text from plain text files"""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()